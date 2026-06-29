"""Convert PROJECT_REPORT.md to a formatted Word document."""
from pathlib import Path
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).parent.parent
MD_PATH = PROJECT_ROOT / "PROJECT_REPORT.md"
OUT_PATH = PROJECT_ROOT / "PROJECT_REPORT.docx"


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def parse_inline(paragraph, text, bold_default=False):
    """Parse **bold**, *italic*, and `code` spans."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            run.bold = bold_default
        chunk = m.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold_default


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_sep(line):
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def build_doc(md_text):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = doc.styles
    if "CodeBlock" not in [s.name for s in styles]:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        rows = table_rows
        table_rows = []
        if len(rows) < 1:
            return
        ncols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                parse_inline(p, cell_text.replace("**", "").strip(), bold_default=(ri == 0))
                if ri == 0:
                    set_cell_shading(cell, "DBEAFE")
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_lines), style="CodeBlock")
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if is_table_sep(stripped):
                i += 1
                continue
            flush_table() if table_rows and not stripped.startswith("|") else None
            table_rows.append(parse_table_row(stripped))
            i += 1
            continue
        else:
            flush_table()

        if stripped in ("---", "***"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:])
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, m.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(stripped.lstrip("> ").strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    flush_table()

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated 2026-06-18 · Tenant Bias Audit · openrouter/owl-alpha")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    return doc


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_doc(md_text)
    doc.save(OUT_PATH)
    print(f"[OK] Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()